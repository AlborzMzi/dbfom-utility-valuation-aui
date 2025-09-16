# src/generate_model.py
# Rebuilds the Excel model and two 1-page documents in /models and /docs.
# Run in Codespaces or locally after installing requirements:
#   pip install -r requirements.txt
#   python src/generate_model.py

import os
from datetime import datetime
import xlsxwriter
from docx import Document
from docx.shared import Inches, Pt

# -------- Assumptions (independent project) --------
BuildMonths = 24
OpsMonths   = 240
EPC         = 62_000_000.0
OMBase      = 325_000.0
Markup      = 0.10
InflAnnual  = 0.02
TownEAR     = 0.0675   # City financing EAR
DebtEAR     = 0.04     # Cost of debt EAR
TaxRate     = 0.265
DebtFrac    = 0.60
EquityFrac  = 0.40

# Derived monthly rates
Town_m = (1 + TownEAR) ** (1/12) - 1
Debt_m = (1 + DebtEAR) ** (1/12) - 1
Infl_m = (1 + InflAnnual) ** (1/12) - 1

def pmt(r, n, pv):
    return r * pv / (1 - (1 + r) ** (-n))

# For summaries
town_payment = pmt(Town_m, OpsMonths, EPC)
debt_payment = pmt(Debt_m, OpsMonths, EPC * DebtFrac)

# Create folders if missing
os.makedirs("models", exist_ok=True)
os.makedirs("docs",   exist_ok=True)

# -------- Excel workbook --------
xl_path = "models/AUI_DBFO_Wastewater_Model.xlsx"
wb = xlsxwriter.Workbook(xl_path)

brand="#1F2937"; accent="#2E6BE6"; zebra="#F4F7FB"
hdr   = wb.add_format({"bold":True,"font_color":"white","bg_color":brand,"border":1,"align":"center","valign":"vcenter"})
title = wb.add_format({"bold":True,"font_size":12,"font_color":accent})
cell  = wb.add_format({"border":1,"text_wrap":True})
wrap  = wb.add_format({"text_wrap":True})
intf  = wb.add_format({"num_format":"0","border":1})
money = wb.add_format({"num_format":"$#,##0.00","border":1})
money0= wb.add_format({"num_format":"$#,##0","border":1})
pct   = wb.add_format({"num_format":"0.00%","border":1})

def prep(ws, tab_color=accent):
    ws.set_tab_color(tab_color); ws.freeze_panes(1,0); ws.set_row(0,22)
    ws.set_footer('&LCreated by Alborz&RPage &P of &N')

# Project_Description
dsc = wb.add_worksheet("Project_Description"); prep(dsc)
dsc.set_column("A:A",120)
dsc.write(0,0,"Aurora Utilities Inc. (AUI) — DBFOM Wastewater Reclamation Facility", title)
dsc.write(2,0,(
  "Independent valuation by Alborz.\n\n"
  "• DBFOM; location: City of Riverview, Alberta\n"
  "• Build 24 months (COD at month 24); Operate 20 years (240 months)\n"
  "• EPC at COD: $62,000,000 (fixed)\n"
  "• City financing: 6.75% EAR (equal monthly payments); only INTEREST is revenue\n"
  "• O&M: cost-plus 10% from $325k/mo base; 2% annual inflation\n"
  "• Capital structure: 60% debt (4.0% EAR) / 40% equity; tax 26.5%\n"
  "• No IDC, no depreciation, no working-capital changes\n\n"
  "Workbook includes Sources&Uses, City receivable, Debt schedule, O&M build, "
  "IS/CF/BS, FCFE/IRR, and a 3×3 sensitivity."
), wrap)

# Executive_Summary
sm = wb.add_worksheet("Executive_Summary"); prep(sm, brand)
sm.set_column("A:A",55); sm.set_column("B:B",28)
sm.write_row("A1", ["Key Item","Value / Notes"], hdr)
sm.write("A3","EPC (capital at COD)", cell);                     sm.write_number("B3", EPC, money0)
sm.write("A4","City level payment (on $62.0M @ 6.75% EAR)", cell); sm.write_number("B4", town_payment, money)
sm.write("A5","Debt level payment (on $37.2M @ 4.00% EAR)", cell); sm.write_number("B5", debt_payment, money)
sm.write("A6","O&M Month-1 revenue (base $325k × 1.10)", cell);    sm.write_number("B6", OMBase*(1+Markup), money)
sm.write("A7","Capital structure", cell);                           sm.write("B7","60% debt / 40% equity", cell)
sm.write("A8","Tax rate", cell);                                    sm.write_number("B8", TaxRate, pct)
sm.write("A10","Decision rule", title); sm.write("B10","Proceed if annualized levered IRR ≥ AUI equity hurdle.", cell)

# Inputs
inp = wb.add_worksheet("Inputs"); prep(inp)
inp.set_column("A:A",52); inp.set_column("B:B",22); inp.set_column("C:C",60)
inp.write_row("A1", ["Parameter","Value","Notes"], hdr)
rows=[("Construction months (BuildMonths)",BuildMonths,"COD at end of month 24"),
      ("Operations months (OpsMonths)",OpsMonths,"20 years after COD"),
      ("EPC cost at COD (EPC)",EPC,"Fixed EPC at COD"),
      ("O&M base monthly cost at COD (OMBase)",OMBase,"At service commencement"),
      ("O&M markup (Markup)",Markup,"10% markup on cost"),
      ("Annual inflation on O&M (InflAnnual)",InflAnnual,"2% / year"),
      ("City financing rate (EAR) (TownEAR)",TownEAR,"Equal monthly payments"),
      ("Cost of debt (EAR) (DebtEAR)",DebtEAR,"Amortized 20 years"),
      ("Effective tax rate (TaxRate)",TaxRate,"Applied when EBT>0"),
      ("Debt fraction (DebtFrac)",DebtFrac,"60% debt"),
      ("Equity fraction (EquityFrac)",EquityFrac,"40% equity")]
for i,(n,v,nt) in enumerate(rows, start=2):
    inp.write(f"A{i}", n, cell)
    inp.write_number(f"B{i}", v, pct if isinstance(v,float) and v<=1 else (money0 if v>=1000 else intf))
    inp.write(f"C{i}", nt, cell)
inp.write("A14","City monthly rate (Town_m)", title); inp.write_formula("B14","=(1+$B$8)^(1/12)-1", pct)
inp.write("A15","Debt monthly rate (Debt_m)", title);  inp.write_formula("B15","=(1+$B$9)^(1/12)-1", pct)
inp.write("A16","Monthly inflation (Infl_m)", title);  inp.write_formula("B16","=(1+$B$7)^(1/12)-1", pct)
inp.write("A18","COD month (CODMonth)", title);        inp.write_formula("B18","=$B$2", intf)
inp.write("A19","Payment months (PaymentMonths)", title); inp.write_formula("B19","=$B$3", intf)

# Named ranges for formulas
for name, ref in [("BuildMonths","=Inputs!$B$2"),("OpsMonths","=Inputs!$B$3"),("EPC","=Inputs!$B$4"),
                  ("OMBase","=Inputs!$B$5"),("Markup","=Inputs!$B$6"),("InflAnnual","=Inputs!$B$7"),
                  ("TownEAR","=Inputs!$B$8"),("DebtEAR","=Inputs!$B$9"),("TaxRate","=Inputs!$B$10"),
                  ("DebtFrac","=Inputs!$B$11"),("EquityFrac","=Inputs!$B$12"),("Town_m","=Inputs!$B$14"),
                  ("Debt_m","=Inputs!$B$15"),("Infl_m","=Inputs!$B$16"),("CODMonth","=Inputs!$B$18"),
                  ("PaymentMonths","=Inputs!$B$19")]:
    wb.define_name(name, ref)

# Sources & Uses
su = wb.add_worksheet("Sources&Uses"); prep(su)
su.set_column("A:A",42); su.set_column("B:B",22)
su.write_row("A1", ["Sources & Uses at COD",""], hdr)
su.write("A3","EPC (construction outlay)", cell); su.write_formula("B3","=EPC", money0)
su.write("A5","Debt draw (60%)", cell);          su.write_formula("B5","=EPC*DebtFrac", money0)
su.write("A6","Equity injection (40%)", cell);   su.write_formula("B6","=EPC*EquityFrac", money0)
su.write("A8","Check: Sources - Uses", cell);    su.write_formula("B8","=B5+B6-B3", money0)

# City_Receivable amortization
tr = wb.add_worksheet("City_Receivable"); prep(tr)
tr.set_column("A:A",10); tr.set_column("B:F",18)
tr.write_row("A1", ["OpMonth","Beg_Receivable","City_Payment","Interest_Revenue","Principal","End_Receivable"], hdr)
for r in range(2, 242):
    tr.write_number(r-1,0,r-1,intf)
    tr.write_formula(r-1,1, "=IF(ROW()=2,EPC,OFFSET(F1,ROW()-3,0))", money)
    tr.write_formula(r-1,2, "=-PMT(Town_m,PaymentMonths,EPC)", money)
    tr.write_formula(r-1,3, f"=B{r}*Town_m", money)
    tr.write_formula(r-1,4, f"=C{r}-D{r}", money)
    tr.write_formula(r-1,5, f"=B{r}-E{r}", money)

# Debt_Schedule
ds = wb.add_worksheet("Debt_Schedule"); prep(ds)
ds.set_column("A:A",10); ds.set_column("B:F",18)
ds.write_row("A1", ["OpMonth","Beg_Debt","Debt_Payment","Interest_Expense","Principal_Repay","End_Debt"], hdr)
for r in range(2, 242):
    ds.write_number(r-1,0,r-1,intf)
    ds.write_formula(r-1,1, "=IF(ROW()=2,EPC*DebtFrac,OFFSET(F1,ROW()-3,0))", money)
    ds.write_formula(r-1,2, "=-PMT(Debt_m,PaymentMonths,EPC*DebtFrac)", money)
    ds.write_formula(r-1,3, f"=B{r}*Debt_m", money)
    ds.write_formula(r-1,4, f"=C{r}-D{r}", money)
    ds.write_formula(r-1,5, f"=B{r}-E{r}", money)

# OM
om = wb.add_worksheet("OM"); prep(om)
om.set_column("A:A",10); om.set_column("B:D",18)
om.write_row("A1", ["OpMonth","OM_Cost","OM_Revenue","OM_Margin"], hdr)
for r in range(2, 242):
    om.write_number(r-1,0,r-1,intf)
    om.write_formula(r-1,1, f"=OMBase*(1+Infl_m)^(A{r}-1)", money)
    om.write_formula(r-1,2, f"=B{r}*(1+Markup)", money)
    om.write_formula(r-1,3, f"=C{r}-B{r}", money)

# IS_Monthly
isf = wb.add_worksheet("IS_Monthly"); prep(isf)
isf.set_column("A:A",8); isf.set_column("B:K",18)
isf.write_row("A1", ["Month","OpIndex","Interest_Revenue","OM_Revenue","Total_Revenue","OM_Cost","EBIT","Debt_Interest_Expense","EBT","Tax","Net_Income"], hdr)
for r in range(2, 266):
    isf.write_number(r-1,0,r-1,intf)
    isf.write_formula(r-1,1, f"=IF(A{r}>BuildMonths,A{r}-BuildMonths,0)", intf)
    isf.write_formula(r-1,2, f"=IF($B{r}>0, INDEX(City_Receivable!$D$2:$D$241,$B{r}), 0)", money)
    isf.write_formula(r-1,3, f"=IF($B{r}>0, INDEX(OM!$C$2:$C$241,$B{r}), 0)", money)
    isf.write_formula(r-1,4, f"=C{r}+D{r}", money)
    isf.write_formula(r-1,5, f"=IF($B{r}>0, INDEX(OM!$B$2:$B$241,$B{r}), 0)", money)
    isf.write_formula(r-1,6, f"=E{r}-F{r}", money)
    isf.write_formula(r-1,7, f"=IF($B{r}>0, INDEX(Debt_Schedule!$D$2:$D$241,$B{r}), 0)", money)
    isf.write_formula(r-1,8, f"=G{r}-H{r}", money)
    isf.write_formula(r-1,9, f"=IF(I{r}>0, I{r}*TaxRate, 0)", money)
    isf.write_formula(r-1,10,f"=I{r}-J{r}", money)

# CF_Monthly
cf = wb.add_worksheet("CF_Monthly"); prep(cf)
cf.set_column("A:A",8); cf.set_column("B:F",18)
cf.write_row("A1", ["Month","CFO","CFI","CFF","Net_Change","Cash_Balance"], hdr)
for r in range(2, 266):
    cf.write_number(r-1,0,r-1,intf)
    cf.write_formula(r-1,1, f"=IS_Monthly!K{r} + IF(IS_Monthly!$B{r}>0, INDEX(City_Receivable!$E$2:$E$241, IS_Monthly!$B{r}), 0)", money)
    cf.write_formula(r-1,2, f"=IF(A{r}=BuildMonths,-EPC,0)", money)
    cf.write_formula(r-1,3, f"=IF(A{r}=BuildMonths,EPC*DebtFrac+EPC*EquityFrac,0) - IF(IS_Monthly!$B{r}>0, INDEX(Debt_Schedule!$E$2:$E$241, IS_Monthly!$B{r}), 0)", money)
    cf.write_formula(r-1,4, f"=B{r}+C{r}+D{r}", money)
    cf.write_formula(r-1,5, f"=IF(ROW()=2,E{r},OFFSET(F1,ROW()-3,0)+E{r})", money)

# BS_Monthly
bs = wb.add_worksheet("BS_Monthly"); prep(bs)
bs.set_column("A:A",8); bs.set_column("B:J",18)
bs.write_row("A1", ["Month","Cash","Receivable","Total_Assets","Debt","Paid_in_Equity","Retained_Earnings","Total_Eq","Liab+Eq","Balance_Check"], hdr)
for r in range(2, 266):
    bs.write_number(r-1,0,r-1,intf)
    bs.write_formula(r-1,1, f"=CF_Monthly!F{r}", money)
    bs.write_formula(r-1,2, f"=IF(A{r}<BuildMonths,0, IF(A{r}=BuildMonths,EPC, IF(IS_Monthly!$B{r}>0, INDEX(City_Receivable!$F$2:$F$241, IS_Monthly!$B{r}), 0)))", money)
    bs.write_formula(r-1,3, f"=B{r}+C{r}", money)
    bs.write_formula(r-1,4, f"=IF(A{r}<BuildMonths,0, IF(A{r}=BuildMonths,EPC*DebtFrac, IF(IS_Monthly!$B{r}>0, INDEX(Debt_Schedule!$F$2:$F$241, IS_Monthly!$B{r}), 0)))", money)
    bs.write_formula(r-1,5, f"=IF(A{r}<BuildMonths,0, EPC*EquityFrac)", money)
    bs.write_formula(r-1,6, f"=IF(ROW()=2,IS_Monthly!K{r}, OFFSET(G1,ROW()-3,0)+IS_Monthly!K{r})", money)
    bs.write_formula(r-1,7, f"=F{r}+G{r}", money)
    bs.write_formula(r-1,8, f"=E{r}+H{r}", money)
    bs.write_formula(r-1,9, f"=D{r}-I{r}", money)

# Levered_DCF_IRR (uses Excel IRR)
irr = wb.add_worksheet("Levered_DCF_IRR"); prep(irr)
irr.set_column("A:A",8); irr.set_column("B:B",22)
irr.write_row("A1", ["Month","Equity_CF"], hdr)
for r in range(2, 266):
    irr.write_number(r-1,0,r-1,intf)
    irr.write_formula(r-1,1, f"=IF(A{r}=BuildMonths,-EPC*EquityFrac, IF(A{r}>BuildMonths, CF_Monthly!B{r} - INDEX(Debt_Schedule!$E$2:$E$241, IS_Monthly!$B{r}), 0))", money)
irr.write("D2","Levered IRR (monthly)", title);   irr.write_formula("E2","=IRR(B2:B265)")
irr.write("D3","Levered IRR (annualized)", title); irr.write_formula("E3","=(1+E2)^12-1")

# Sensitivity (base cell link for now)
sens = wb.add_worksheet("Sensitivity_3x3"); prep(sens)
sens.set_column("A:A",28); sens.set_column("B:D",14)
sens.write_row("A1", ["Levered IRR (annualized) — Sensitivity Grid", "", "", ""], hdr)
sens.write("A3","Markup ↓ / City EAR →", cell)
sens.write_number("B4",0.0575,pct); sens.write_number("C4",0.0675,pct); sens.write_number("D4",0.0775,pct)
sens.write_number("A5",0.05,pct);   sens.write_number("A6",0.10,pct);   sens.write_number("A7",0.15,pct)
sens.write_formula("B6","=Levered_DCF_IRR!E3")  # base-case pointer; expand later if desired

wb.close()

# -------- Word one-pagers --------
def one_pager(path, title_text, body_paragraphs):
    d = Document()
    for s in d.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6); s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    r = d.add_paragraph().add_run(title_text); r.bold=True; r.font.size=Pt(14)
    d.add_paragraph(f"Prepared by: Alborz    |    Date: {datetime.now():%Y-%m-%d}").italic=True
    for p in body_paragraphs:
        d.add_paragraph(p)
    d.save(path)

one_pager("docs/Project_Description.docx",
          "Aurora Utilities Inc. (AUI) — Project Description",
          ["DBFOM wastewater facility (Riverview, AB). Build 24 months; operate 20 years. EPC at COD $62MM.",
           "City repays via equal monthly payments at 6.75% EAR; only interest is recognized as revenue.",
           "O&M billed cost-plus 10% from $325k/mo base, indexed 2%/yr.",
           "Capital structure 60% debt (4.0% EAR) / 40% equity; tax 26.5%; no IDC/DEP/WC.",
           "Workbook: Sources&Uses, City receivable, Debt schedule, O&M build, IS/CF/BS, FCFE/IRR, sensitivity."])

one_pager("docs/Executive_Summary.docx",
          "Executive Summary — AUI DBFOM Wastewater Reclamation Facility",
          [f"City level payment (on $62.0MM @ 6.75% EAR, 240 mo): ${town_payment:,.2f}",
           f"Debt level payment (on $37.2MM @ 4.00% EAR, 240 mo): ${debt_payment:,.2f}",
           f"Month-1 O&M revenue (base $325k × 1.10): ${(OMBase*(1+Markup)):,.2f}",
           "Decision rule: proceed if annualized levered IRR ≥ AUI equity hurdle."])

print("Built:", xl_path, "and docs in /docs")
